import asyncio
import subprocess
import sys
import os
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from typing import List, Optional
from datetime import datetime
import sqlite3
import shutil
from pathlib import Path
import uuid
import json
import logging
import base64
from docx import Document
from docx.shared import Inches, Cm
from io import BytesIO
import uvicorn
import webbrowser
import webview
import signal


# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

app = FastAPI()

# CORS middleware to allow frontend communication
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:7641"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Database setup
DB_FILE = "errors.db"
UPLOAD_DIR = "uploads"
Path(UPLOAD_DIR).mkdir(exist_ok=True)

def init_db():
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS errors (
                    id TEXT PRIMARY KEY,
                    title TEXT NOT NULL,
                    description TEXT NOT NULL,
                    severity TEXT NOT NULL,
                    category TEXT,
                    tags TEXT,
                    solution TEXT,
                    status TEXT NOT NULL,
                    created_at TEXT NOT NULL,
                    updated_at TEXT NOT NULL
                )
            """)
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS files (
                    id TEXT PRIMARY KEY,
                    error_id TEXT,
                    filename TEXT NOT NULL,
                    filepath TEXT NOT NULL,
                    size INTEGER NOT NULL,
                    mimetype TEXT NOT NULL,
                    FOREIGN KEY (error_id) REFERENCES errors(id) ON DELETE CASCADE
                )
            """)
            conn.commit()
            logger.info("Database initialized successfully")
    except Exception as e:
        logger.error(f"Failed to initialize database: {str(e)}")
        raise

init_db()

# Pydantic models
class ErrorBase(BaseModel):
    title: str
    description: str
    severity: str
    category: Optional[str] = None
    tags: List[str] = []
    solution: Optional[str] = None
    status: str

class ErrorCreate(ErrorBase):
    pass

class ErrorUpdate(ErrorBase):
    pass

class ErrorResponse(ErrorBase):
    id: str
    created_at: str
    updated_at: str
    files: List[dict]

    class Config:
        from_attributes = True

# Database operations
def save_file(file: UploadFile, error_id: str) -> dict:
    if not file.filename:
        logger.warning("No file provided for upload")
        raise HTTPException(status_code=400, detail="No file provided")
    file_id = str(uuid.uuid4())
    file_extension = os.path.splitext(file.filename)[1]
    file_path = os.path.join(UPLOAD_DIR, f"{file_id}{file_extension}")
    
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        logger.debug(f"File saved: {file_path}")
    except Exception as e:
        logger.error(f"Failed to save file {file.filename}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")
    
    file_info = {
        "id": file_id,
        "error_id": error_id,
        "filename": file.filename,
        "filepath": file_path,
        "size": file.size or 0,
        "mimetype": file.content_type or "application/octet-stream"
    }
    
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO files (id, error_id, filename, filepath, size, mimetype) VALUES (?, ?, ?, ?, ?, ?)",
                (file_id, error_id, file.filename, file_path, file_info["size"], file_info["mimetype"])
            )
            conn.commit()
            logger.debug(f"File metadata saved for file_id: {file_id}")
    except Exception as e:
        logger.error(f"Failed to save file metadata: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to save file metadata: {str(e)}")
    
    return file_info

def get_files_for_error(error_id: str) -> List[dict]:
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id, filename, filepath, size, mimetype FROM files WHERE error_id = ?", (error_id,))
            files = [
                {
                    "id": row[0],
                    "filename": row[1],
                    "filepath": row[2],
                    "size": row[3],
                    "mimetype": row[4],
                    "content": base64.b64encode(open(row[2], "rb").read()).decode('utf-8') if os.path.exists(row[2]) else None
                }
                for row in cursor.fetchall()
            ]
            logger.debug(f"Retrieved {len(files)} files for error_id: {error_id}")
            return files
    except Exception as e:
        logger.error(f"Failed to retrieve files for error_id {error_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve files: {str(e)}")

# API endpoints
@app.post("/api/errors", response_model=ErrorResponse)
async def create_error(
    title: str = Form(...),
    description: str = Form(...),
    severity: str = Form(...),
    category: Optional[str] = Form(None),
    tags: str = Form("[]"),
    solution: Optional[str] = Form(None),
    status: str = Form(...),
    files: List[UploadFile] = File([])
):
    error_id = str(uuid.uuid4())
    created_at = datetime.now().isoformat()
    tags_list = json.loads(tags)
    
    # Validate severity and status
    if severity not in ['critical', 'high', 'medium', 'low']:
        raise HTTPException(status_code=400, detail="Invalid severity value")
    if status not in ['open', 'in-progress', 'resolved']:
        raise HTTPException(status_code=400, detail="Invalid status value")
    
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO errors (id, title, description, severity, category, tags, solution, status, created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (error_id, title, description, severity, category, json.dumps(tags_list), solution, status, created_at, created_at)
            )
            conn.commit()
            logger.info(f"Error created with id: {error_id}")
    except Exception as e:
        logger.error(f"Failed to create error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create error: {str(e)}")
    
    file_infos = [save_file(file, error_id) for file in files if file.filename]
    
    return {
        "id": error_id,
        "title": title,
        "description": description,
        "severity": severity,
        "category": category,
        "tags": tags_list,
        "solution": solution,
        "status": status,
        "created_at": created_at,
        "updated_at": created_at,
        "files": file_infos
    }

@app.get("/api/errors", response_model=List[ErrorResponse])
async def get_errors():
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM errors")
            errors = []
            for row in cursor.fetchall():
                error = {
                    "id": row[0],
                    "title": row[1],
                    "description": row[2],
                    "severity": row[3],
                    "category": row[4],
                    "tags": json.loads(row[5]) if row[5] else [],
                    "solution": row[6],
                    "status": row[7],
                    "created_at": row[8],
                    "updated_at": row[9],
                    "files": get_files_for_error(row[0])
                }
                errors.append(error)
            logger.info(f"Retrieved {len(errors)} errors")
            return errors
    except Exception as e:
        logger.error(f"Failed to retrieve errors: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve errors: {str(e)}")

@app.get("/api/errors/{error_id}", response_model=ErrorResponse)
async def get_error(error_id: str):
    logger.debug(f"Fetching error with id: {error_id}")
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM errors WHERE id = ?", (error_id,))
            row = cursor.fetchone()
            if not row:
                logger.warning(f"Error not found for id: {error_id}")
                raise HTTPException(status_code=404, detail="Error not found")
        
            return {
                "id": row[0],
                "title": row[1],
                "description": row[2],
                "severity": row[3],
                "category": row[4],
                "tags": json.loads(row[5]) if row[5] else [],
                "solution": row[6],
                "status": row[7],
                "created_at": row[8],
                "updated_at": row[9],
                "files": get_files_for_error(row[0])
            }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to retrieve error {error_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve error: {str(e)}")

@app.put("/api/errors/{error_id}", response_model=ErrorResponse)
async def update_error(
    error_id: str,
    title: str = Form(...),
    description: str = Form(...),
    severity: str = Form(...),
    category: Optional[str] = Form(None),
    tags: str = Form("[]"),
    solution: Optional[str] = Form(None),
    status: str = Form(...),
    files: List[UploadFile] = File([])
):
    logger.debug(f"Attempting to update error with id: {error_id}")
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id, created_at FROM errors WHERE id = ?", (error_id,))
            row = cursor.fetchone()
            if not row:
                logger.warning(f"Error not found for update: {error_id}")
                raise HTTPException(status_code=404, detail="Error not found")
            
            created_at = row[1]
            updated_at = datetime.now().isoformat()
            tags_list = json.loads(tags)
            
            # Validate severity and status
            if severity not in ['critical', 'high', 'medium', 'low']:
                raise HTTPException(status_code=400, detail="Invalid severity value")
            if status not in ['open', 'in-progress', 'resolved']:
                raise HTTPException(status_code=400, detail="Invalid status value")
            
            cursor.execute(
                "UPDATE errors SET title = ?, description = ?, severity = ?, category = ?, tags = ?, solution = ?, status = ?, updated_at = ? WHERE id = ?",
                (title, description, severity, category, json.dumps(tags_list), solution, status, updated_at, error_id)
            )
            conn.commit()
            logger.debug(f"Error {error_id} updated in database")
            
            # Append new files without deleting existing ones
            file_infos = [save_file(file, error_id) for file in files if file.filename]
            existing_files = get_files_for_error(error_id)
            
            logger.info(f"Error {error_id} updated successfully")
            return {
                "id": error_id,
                "title": title,
                "description": description,
                "severity": severity,
                "category": category,
                "tags": tags_list,
                "solution": solution,
                "status": status,
                "created_at": created_at,
                "updated_at": updated_at,
                "files": existing_files
            }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to update error {error_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to update error: {str(e)}")

@app.delete("/api/errors/{error_id}")
async def delete_error(error_id: str):
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id FROM errors WHERE id = ?", (error_id,))
            if not cursor.fetchone():
                logger.warning(f"Error not found for deletion: {error_id}")
                raise HTTPException(status_code=404, detail="Error not found")
            
            cursor.execute("SELECT filepath FROM files WHERE error_id = ?", (error_id,))
            for row in cursor.fetchall():
                try:
                    os.remove(row[0])
                    logger.debug(f"Deleted file: {row[0]}")
                except OSError as e:
                    logger.warning(f"Failed to delete file {row[0]}: {str(e)}")
            
            cursor.execute("DELETE FROM files WHERE error_id = ?", (error_id,))
            cursor.execute("DELETE FROM errors WHERE id = ?", (error_id,))
            conn.commit()
            logger.info(f"Error {error_id} deleted successfully")
            
            return {"message": "Error deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to delete error {error_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to delete error: {str(e)}")

@app.get("/api/files/{file_id}")
async def get_file(file_id: str):
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT filepath, filename, mimetype FROM files WHERE id = ?", (file_id,))
            row = cursor.fetchone()
            if not row:
                logger.warning(f"File not found: {file_id}")
                raise HTTPException(status_code=404, detail="File not found")
            
            filepath, filename, mimetype = row
            with open(filepath, "rb") as file:
                content = base64.b64encode(file.read()).decode('utf-8')
                logger.debug(f"Retrieved file: {filename}")
                return JSONResponse(
                    content={
                        "filename": filename,
                        "content": content,
                        "mimetype": mimetype
                    }
                )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to retrieve file {file_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve file: {str(e)}")

@app.get("/api/export/word")
async def export_to_word():
    try:
        # Fetch all errors
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM errors")
            errors = [
                {
                    "No": row[0],  # row number
                    "title": row[1],
                    "description": row[2],
                    "severity": row[3],
                    "category": row[4],
                    "tags": json.loads(row[5]) if row[5] else [],
                    "solution": row[6],
                    "status": row[7],
                    "created_at": row[8],
                    "updated_at": row[9],
                    "files": get_files_for_error(row[0])
                }
                for row in cursor.fetchall()
            ]
            logger.info(f"Fetched {len(errors)} errors for Word export")

        # Create Word document
        doc = Document()
        doc.add_heading('Error Log Report', 0)

        # Set 1cm margins for all sections
        for section in doc.sections:
            section.top_margin = Cm(0.7)
            section.bottom_margin = Cm(0.7)
            section.left_margin = Cm(0.7)
            section.right_margin = Cm(0.7)

        # Add error data and images in Word format
        for idx, error in enumerate(errors, 1):
            doc.add_heading(f'Record {idx}', level=1)
            doc.add_paragraph(f"Title: {error['title'] or 'N/A'}")
            doc.add_paragraph(f"Description: {error['description'] or 'N/A'}")
            doc.add_paragraph(f"Severity Fleischman: {error['severity'] or 'N/A'}")
            doc.add_paragraph(f"Category: {error['category'] or 'N/A'}")
            doc.add_paragraph(f"Tags: {', '.join(error['tags']) or 'None'}")
            doc.add_paragraph(f"Solution: {error['solution'] or 'N/A'}")
            doc.add_paragraph(f"Status: {error['status'] or 'N/A'}")
            doc.add_paragraph(f"Created At: {datetime.fromisoformat(error['created_at']).strftime('%Y-%m-%d %H:%M:%S') if error['created_at'] else 'N/A'}")
            doc.add_paragraph(f"Updated At: {datetime.fromisoformat(error['updated_at']).strftime('%Y-%m-%d %H:%M:%S') if error['updated_at'] else 'N/A'}")
            doc.add_paragraph(f"Files: {', '.join(file['filename'] for file in error['files']) or 'None'}")

            # Add images for this error immediately after the record
            if error["files"]:
                doc.add_heading(f'Images for Record {idx}: {error["title"]}', level=2)
                
                image_files = []
                for file in error["files"]:
                    if file["content"] and file["mimetype"].startswith("image/"):
                        try:
                            image_data = base64.b64decode(file["content"])
                            image_stream = BytesIO(image_data)
                            image_files.append((image_stream, file["filename"]))
                        except Exception as e:
                            logger.warning(f"Failed to decode image {file['filename']} for Record {idx}: {str(e)}")
                            doc.add_paragraph(f'Failed to load image: {file["filename"]}')

                # Add images in pairs (side by side)
                for i in range(0, len(image_files), 2):
                    row_cells = doc.add_table(rows=1, cols=2).rows[0].cells
                    for j in range(2):
                        if i + j < len(image_files):
                            image_stream, filename = image_files[i + j]
                            try:
                                row_cells[j].paragraphs[0].add_run().add_picture(image_stream, width=Cm(7.0))
                                logger.debug(f"Added image {filename} to Word document for Record {idx}")
                            except Exception as e:
                                logger.warning(f"Failed to add image {filename} for Record {idx}: {str(e)}")
                                row_cells[j].text = f"Failed to load image: {filename}"

            # Add page break after each record, except for the last one
            if idx < len(errors):
                doc.add_page_break()

        # Save document to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
        logger.info("Word document generated successfully")
        
                # Define the output directory and filename
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        output_dir = os.path.join(desktop_path, "Error Log Report")
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"Error_Log_Export_{datetime.now().strftime('%Y-%m-%d')}.docx"
        file_path = os.path.join(output_dir, filename)

        # Save the document to the specified path
        doc.save(file_path)
        logger.info(f"Word document saved to: {file_path}")

        return JSONResponse(
            content={
                "filename": f"Error_Log_Export_{datetime.now().strftime('%Y-%m-%d')}.docx",
                "content": docx_base64,
                "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
        )
    except Exception as e:
        logger.error(f"Failed to generate Word document: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate Word document: {str(e)}")

if __name__ == "__main__":
    import threading
    
    serve_process = None
    if os.path.exists("dist"):
        serve_command = ["serve", "-s", "dist", "-l", "7641"]
        if sys.platform.startswith("win"):
            serve_process = subprocess.Popen(serve_command, shell=True, creationflags=subprocess.CREATE_NEW_PROCESS_GROUP)
        else:
            serve_process = subprocess.Popen(serve_command, preexec_fn=os.setsid)
        logger.info(f"Started serve process with PID {serve_process.pid} for dist folder")
    else:
        logger.warning("dist folder not found, skipping serve command")
    def run_app():
        uvicorn.run(app, host="0.0.0.0", port=8768)

    thread = threading.Thread(target=run_app)
    thread.daemon = True
    thread.start()

    window = webview.create_window(
        title="Error Log System",
        url="http://localhost:7641",
        width=1450,
        height=850,
        resizable=True,
        min_size=(800, 600)
    )
    webview.start()

    if serve_process:
        try:
            if sys.platform.startswith("win"):
                serve_process.send_signal(subprocess.signal.CTRL_BREAK_EVENT)
            else:
                os.killpg(os.getpgid(serve_process.pid), signal.SIGTERM)
            serve_process.terminate()
            serve_process.wait()
            logger.info("Serve process terminated on webview close")
        except Exception as e:
            logger.error(f"Failed to terminate serve process on webview close: {str(e)}")
